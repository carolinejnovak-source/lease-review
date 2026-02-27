"""
Apply tracked changes (redlines) and comment annotations to a .docx file.
Produces real attorney-style tracked changes (w:ins/w:del) plus highlighted
comment annotations for issues that can't be redlined directly.
"""
import re
from copy import deepcopy
from datetime import datetime, timezone
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

AUTHOR = "VIP Medical Legal"
DATE = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


# ── Text helpers ─────────────────────────────────────────────────────────────

def _normalize(text):
    """Collapse whitespace for fuzzy matching."""
    return re.sub(r'\s+', ' ', text).strip()


# ── XML helpers ──────────────────────────────────────────────────────────────

def _make_run(text, rpr=None):
    r = OxmlElement("w:r")
    if rpr is not None:
        r.append(deepcopy(rpr))
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r


def _make_del(text, rpr, change_id):
    del_el = OxmlElement("w:del")
    del_el.set(qn("w:id"), str(change_id))
    del_el.set(qn("w:author"), AUTHOR)
    del_el.set(qn("w:date"), DATE)
    r = OxmlElement("w:r")
    if rpr is not None:
        r.append(deepcopy(rpr))
    dt = OxmlElement("w:delText")
    dt.set(qn("xml:space"), "preserve")
    dt.text = text
    r.append(dt)
    del_el.append(r)
    return del_el


def _make_ins(text, rpr, change_id):
    ins_el = OxmlElement("w:ins")
    ins_el.set(qn("w:id"), str(change_id))
    ins_el.set(qn("w:author"), AUTHOR)
    ins_el.set(qn("w:date"), DATE)
    r = OxmlElement("w:r")
    if rpr is not None:
        r.append(deepcopy(rpr))
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    ins_el.append(r)
    return ins_el


def _get_rpr(para):
    for run in para.runs:
        rpr = run._r.find(qn("w:rPr"))
        if rpr is not None:
            return deepcopy(rpr)
    return None


# ── Apply a single redline to one paragraph ──────────────────────────────────

def _apply_to_para(para, find_text, replace_text, change_id):
    """
    Try exact then normalized-whitespace match.
    Returns (new_change_id, found_bool).
    """
    full_text = "".join(run.text for run in para.runs)

    actual_find = find_text
    actual_full = full_text

    if find_text not in full_text:
        norm_full = _normalize(full_text)
        norm_find = _normalize(find_text)
        if norm_find in norm_full:
            actual_find = norm_find
            actual_full = norm_full
        else:
            return change_id, False

    rpr = _get_rpr(para)
    idx = actual_full.index(actual_find)
    before = actual_full[:idx]
    after  = actual_full[idx + len(actual_find):]

    pPr = para._p.find(qn("w:pPr"))
    for child in list(para._p):
        if child.tag != qn("w:pPr"):
            para._p.remove(child)
    if pPr is not None and para._p.find(qn("w:pPr")) is None:
        para._p.insert(0, pPr)

    if before:
        para._p.append(_make_run(before, rpr))
    para._p.append(_make_del(actual_find, rpr, change_id))
    change_id += 1
    para._p.append(_make_ins(replace_text, rpr, change_id))
    change_id += 1
    if after:
        para._p.append(_make_run(after, rpr))

    return change_id, True


# ── Comment annotation ────────────────────────────────────────────────────────

def _comment_confidence(para_text: str, section_name: str, keywords: list) -> float:
    """
    Score 0.0–1.0: how confident are we that this paragraph is the right
    anchor for a comment about `section_name`?

    Components:
      +0.40  section name (or most of it) appears verbatim in the paragraph
      +0.30  ALL keywords present
      +0.10  any keyword present (if not all)
      +0.30  paragraph is heading-like (≤ 80 chars)
      +0.15  paragraph is short (≤ 160 chars)
    """
    t = para_text.lower().strip()
    sn = section_name.lower()
    score = 0.0

    if sn in t:
        score += 0.40
    if keywords and all(kw in t for kw in keywords):
        score += 0.30
    elif keywords and any(kw in t for kw in keywords):
        score += 0.10

    length = len(para_text.strip())
    if length <= 80:
        score += 0.30
    elif length <= 160:
        score += 0.15

    return min(score, 1.0)


CONFIDENCE_THRESHOLD = 0.80   # below this → append to end of document


def _insert_comment_annotation(doc, section_name, issue_text, vip_standard="",
                               target_para_idx: int = None) -> int:
    """
    Insert a yellow-highlighted VIP LEGAL REVIEW annotation.

    Strategy:
      - If target_para_idx is provided, insert directly above that paragraph
        (used when a redline was already applied there — guarantees comment
        sits immediately above the redlined text).
      - Otherwise, score every paragraph and insert above the best match.
      - If best score < CONFIDENCE_THRESHOLD → append to bottom of document.

    Returns the paragraph index used (or 999999 when appended to bottom).
    """
    best_para     = None
    best_para_idx = 999999
    best_score    = 0.0

    if target_para_idx is not None and 0 <= target_para_idx < len(doc.paragraphs):
        # Caller already knows the exact paragraph — skip keyword search
        best_para     = doc.paragraphs[target_para_idx]
        best_para_idx = target_para_idx
        best_score    = 1.0   # treat as perfect confidence
    else:
        keywords = [w for w in section_name.lower().split() if len(w) > 3]
        for para_idx, para in enumerate(doc.paragraphs):
            if not para.text.strip():
                continue
            score = _comment_confidence(para.text, section_name, keywords)
            if score > best_score:
                best_score    = score
                best_para     = para
                best_para_idx = para_idx

    # Build the comment paragraph element
    comment_para = OxmlElement('w:p')

    # Paragraph properties: yellow background
    pPr = OxmlElement('w:pPr')
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'FFFF99')
    pPr.append(shd)
    comment_para.append(pPr)

    # Run: bold, red, yellow highlight
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    bold = OxmlElement('w:b')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), 'CC0000')
    highlight = OxmlElement('w:highlight')
    highlight.set(qn('w:val'), 'yellow')
    shd2 = OxmlElement('w:shd')
    shd2.set(qn('w:val'), 'clear')
    shd2.set(qn('w:color'), 'auto')
    shd2.set(qn('w:fill'), 'FFFF99')
    rPr.extend([bold, color, highlight, shd2])
    r.append(rPr)

    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    text = f'⚑ VIP LEGAL REVIEW \u2014 {section_name.upper()}: {issue_text}'
    if vip_standard:
        text += f'  |  VIP STANDARD: {vip_standard}'
    t.text = text
    r.append(t)
    comment_para.append(r)

    # High confidence → insert ABOVE the found paragraph
    # Low confidence  → append to bottom of document
    if best_score >= CONFIDENCE_THRESHOLD and best_para is not None:
        best_para._p.addprevious(comment_para)
        return best_para_idx
    else:
        doc.element.body.append(comment_para)
        return 999999


# ── Public API ───────────────────────────────────────────────────────────────

def apply_redlines(input_path: str, redlines: list, output_path: str,
                   issues: list = None, additional_issues: list = None) -> dict:
    """
    Apply tracked-change redlines to a .docx.
    For redlines where the find-text can't be located, insert a comment annotation.
    For High/Medium issues with no associated redline, also insert comment annotations.

    Returns: {applied, comments, skipped, section_actions}
      section_actions: {section_name: "redline"|"comment"}
    """
    doc = Document(input_path)
    change_id = 1
    applied = 0
    comments = 0
    section_actions = {}   # section -> "redline" | "comment" | "both"
    find_to_section = {}  # find_text (lowered) -> section name, for post-scan

    def _record_action(sec, action):
        existing = section_actions.get(sec)
        if existing is None:
            section_actions[sec] = action
        elif existing != action:
            section_actions[sec] = "both"   # both redline + comment applied
    section_positions = {} # section -> paragraph index in document (for lease-order sort)

    # Index issues by section for quick lookup
    issues_by_section = {}
    for item in (issues or []):
        sec = item.get('section', '')
        if sec:
            issues_by_section[sec] = item

    # ── Pre-scan original doc: find-text positions (before any modifications) ─
    # Lock in a position for every section that has a redline, based on where
    # the find-text appears in the ORIGINAL document.  This gives accurate sort
    # positions even when the redline application later fails (text mismatch).
    for redline_item in redlines:
        find_text = (redline_item.get("find") or "").strip().lower()
        sec       = redline_item.get("section", "")
        if not find_text or not sec:
            continue
        find_norm = re.sub(r'\s+', ' ', find_text).strip()
        for para_idx, para in enumerate(doc.paragraphs):
            para_full = re.sub(r'\s+', ' ', ''.join(para._p.itertext()).lower())
            if find_norm in para_full:
                if para_idx < section_positions.get(sec, 999999):
                    section_positions[sec] = para_idx
                break

    # ── Keyword fallback: approximate positions for still-unpositioned sections ─
    # Build a ranked list of candidate search strings for each section, then
    # try them in order against the original document.  Multiple passes with
    # decreasing specificity so we always find something better than 999999.
    import re as _re

    # Collect redline find-texts per section (most reliable — verbatim lease text)
    section_find_texts: dict = {}
    for r in redlines:
        sec = r.get("section", "")
        find = (r.get("find") or "").strip()
        if sec and find:
            section_find_texts.setdefault(sec, []).append(find.lower())

    # Collect lease_says snippets + quoted substrings per section
    section_lease_says: dict = {}
    for item in (issues or []):
        sec = item.get("section", "")
        ls  = (item.get("lease_says") or "").strip()
        if not sec or not ls:
            continue
        snippets = [ls[:60].lower()]
        # Extract anything in single/double/curly quotes (e.g. 'VIP Medical Group')
        quoted = _re.findall(r"['\u2018\u2019\u201c\u201d\"]([^'\"]{4,60})['\u2018\u2019\u201c\u201d\"]", ls)
        snippets += [q.lower() for q in quoted if len(q) > 4]
        section_lease_says[sec] = snippets

    all_section_names = set(r.get("section", "") for r in redlines) | \
                        set(i.get("section", "") for i in (issues or []))

    for sec in all_section_names:
        if not sec or sec in section_positions:
            continue   # already positioned by find-text pre-scan

        kws = [w for w in sec.lower().split() if len(w) > 3]

        def _para_full_text(para):
            """All text in a paragraph including hyperlinks, bookmarks, etc."""
            return re.sub(r'\s+', ' ', ''.join(para._p.itertext()).lower())

        def _first_match(candidates):
            for cand in candidates:
                if not cand or len(cand) < 4:
                    continue
                cand_norm = re.sub(r'\s+', ' ', cand.lower()).strip()
                for pi, para in enumerate(doc.paragraphs):
                    if cand_norm in _para_full_text(para):
                        return pi
            return None

        # Pass 1: redline find-text (verbatim lease text, highest accuracy)
        pos = _first_match(section_find_texts.get(sec, []))

        # Pass 2: lease_says snippets + quoted substrings
        if pos is None:
            pos = _first_match(section_lease_says.get(sec, []))

        # Pass 3: all section-name keywords in same paragraph
        if pos is None and kws:
            for pi, para in enumerate(doc.paragraphs):
                if all(kw in para.text.lower() for kw in kws):
                    pos = pi
                    break

        # Pass 4: longest single keyword
        if pos is None and kws:
            pos = _first_match([max(kws, key=len)])

        if pos is not None:
            section_positions[sec] = pos

    # ── Step 1: Apply redlines ────────────────────────────────────────────────
    for redline in redlines:
        find    = (redline.get("find") or "").strip()
        replace = (redline.get("replace") or "").strip()
        reason  = redline.get("reason", "")
        section = redline.get("section", "")

        if not find or not replace or find == replace:
            continue

        if section and find:
            find_to_section[find.lower()[:80]] = section  # for post-scan mapping

        found = False

        # Search paragraphs — apply to ALL occurrences (not just the first)
        # so repeated clauses (e.g. "June 1, 2026" or wrong entity name throughout)
        # are all corrected in one pass.
        for para_idx, para in enumerate(doc.paragraphs):
            change_id, ok = _apply_to_para(para, find, replace, change_id)
            if ok:
                applied += 1
                _record_action(section, "redline")
                # Record only the FIRST (earliest) occurrence for sort-order purposes
                if not found:
                    section_positions[section] = para_idx
                found = True
                # Do NOT break — continue to replace all remaining occurrences

        # Search table cells — also apply to all occurrences
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        change_id, ok = _apply_to_para(para, find, replace, change_id)
                        if ok:
                            found = True
                            applied += 1
                            _record_action(section, "redline")
                            if section_positions.get(section, 999999) >= 999998:
                                section_positions[section] = 999998

        # Text not found → fall back to comment annotation
        if not found and section_actions.get(section) not in ("redline", "comment", "both"):
            issue_obj  = issues_by_section.get(section, {})
            issue_text = issue_obj.get('issue') or reason or "Review required per VIP standards"
            vip_std    = issue_obj.get('vip_standard') or ""
            para_idx = _insert_comment_annotation(doc, section, issue_text, vip_std)
            _record_action(section, "comment")
            # Keep earliest position — never overwrite a lower (earlier) value
            if para_idx < section_positions.get(section, 999999):
                section_positions[section] = para_idx
            comments += 1

    # ── Step 2: Comments for High/Medium issues with no comment yet ───────────
    if issues:
        for issue in issues:
            sec      = issue.get('section', '')
            priority = issue.get('priority', 'Low')
            status   = issue.get('status', 'pass')

            if status == 'pass' or priority not in ('High', 'Medium'):
                continue
            # Skip if we already inserted a comment for this section
            if section_actions.get(sec) in ("comment", "both"):
                continue

            issue_text = issue.get('issue') or 'Review required per VIP standards'
            vip_std    = issue.get('vip_standard') or ''

            # If a redline was applied for this section, we know exactly which
            # paragraph it sits in — insert the comment immediately above it.
            known_pos = section_positions.get(sec)
            target_idx = known_pos if (known_pos is not None and known_pos < 999990
                                       and section_actions.get(sec) == "redline") else None

            para_idx = _insert_comment_annotation(doc, sec, issue_text, vip_std,
                                                  target_para_idx=target_idx)
            _record_action(sec, "comment")
            # Keep earliest position — never overwrite a lower (earlier) value
            if para_idx < section_positions.get(sec, 999999):
                section_positions[sec] = para_idx
            comments += 1

    # (Passing items have no redline/comment, so no position needed —
    #  they sort to the bottom of the lease-order view with position 999999.)

    # ── Additional Issues: append non-checklist red flags at end of document ──
    if additional_issues:
        hdr = OxmlElement('w:p')
        hdr_r = OxmlElement('w:r')
        hdr_rpr = OxmlElement('w:rPr')
        hdr_b = OxmlElement('w:b')
        hdr_rpr.append(hdr_b)
        hdr_r.append(hdr_rpr)
        hdr_t = OxmlElement('w:t')
        hdr_t.text = '\u2756 ADDITIONAL VIP REVIEW ITEMS \u2014 NON-STANDARD CLAUSES'
        hdr_r.append(hdr_t)
        hdr.append(hdr_r)
        doc.element.body.append(hdr)

        for ai_item in additional_issues:
            title   = (ai_item.get('title') or 'Additional Issue').upper()
            concern = ai_item.get('concern') or ''
            action  = ai_item.get('suggested_action') or ''
            ai_para = OxmlElement('w:p')
            ai_ppr  = OxmlElement('w:pPr')
            ai_shd  = OxmlElement('w:shd')
            ai_shd.set(qn('w:val'), 'clear')
            ai_shd.set(qn('w:color'), 'auto')
            ai_shd.set(qn('w:fill'), 'FFF3CD')
            ai_ppr.append(ai_shd)
            ai_para.append(ai_ppr)
            ai_r   = OxmlElement('w:r')
            ai_rpr = OxmlElement('w:rPr')
            ai_b2  = OxmlElement('w:b')
            ai_rpr.append(ai_b2)
            ai_clr = OxmlElement('w:color')
            ai_clr.set(qn('w:val'), 'B45309')
            ai_rpr.append(ai_clr)
            ai_r.append(ai_rpr)
            ai_t = OxmlElement('w:t')
            ai_t.text = f'\u2691 {title}: {concern}'
            if action:
                ai_t.text += f'  |  ACTION: {action}'
            ai_r.append(ai_t)
            ai_para.append(ai_r)
            doc.element.body.append(ai_para)

    doc.save(output_path)

    # ── Final re-scan: true document order from the saved file ───────────────
    # Comment insertions (addprevious) shift paragraph indices, so we re-open
    # the saved file and read the actual positions of every annotation/redline.
    try:
        import sys
        final_doc = Document(output_path)
        final_positions = {}   # fresh — scan order = true document order

        MARKER = '\u2691 VIP LEGAL REVIEW \u2014 '

        # Build a lowercase lookup: section_name.upper() -> original section name
        # so we can match "ESTOPPEL CERTIFICATES - SECTION 15.3" → "Estoppel Certificates"
        sec_upper_map = {s.upper(): s for s in section_actions}

        for para_idx, para in enumerate(final_doc.paragraphs):
            para_text = para.text  # concatenates all w:t runs

            # 1. Comment markers — "⚑ VIP LEGAL REVIEW — SECTIONNAME: …"
            #    Section names may have suffixes like " - SECTION 15.3"; use
            #    startswith matching against every known section.
            if MARKER in para_text:
                after = para_text[para_text.index(MARKER) + len(MARKER):]
                if ':' in after:
                    raw = after[:after.index(':')].strip().upper()
                    # Exact match first, then startswith
                    matched_sec = sec_upper_map.get(raw)
                    if matched_sec is None:
                        for known_upper, known_sec in sec_upper_map.items():
                            if raw.startswith(known_upper):
                                matched_sec = known_sec
                                break
                    if matched_sec and para_idx < final_positions.get(matched_sec, 999999):
                        final_positions[matched_sec] = para_idx

            # 2. Redline paragraphs — check w:delText (tracked deletions)
            del_texts = [
                el.text or '' for el in para._p.iter()
                if el.tag == qn('w:delText')
            ]
            if del_texts:
                combined = ''.join(del_texts).lower()
                for find_key, sec in find_to_section.items():
                    if find_key in combined:
                        if para_idx < final_positions.get(sec, 999999):
                            final_positions[sec] = para_idx
                        break

        # Merge: keep the EARLIEST (lowest) position from either source.
        # section_positions holds pre-scan / application positions (original doc).
        # final_positions holds re-scan positions (saved doc, shifted by insertions).
        # Taking the minimum ensures a bottom-dump comment (pos ~460) never
        # overrides a correct pre-scan position (pos 2) for the same section.
        for sec, pos in section_positions.items():
            if pos < final_positions.get(sec, 999999):
                final_positions[sec] = pos
        section_positions = final_positions
        print(f"[rescan] positions: {sorted(section_positions.items(), key=lambda x:x[1])[:8]}", file=sys.stderr)

    except Exception as e:
        import sys, traceback
        print(f"[rescan ERROR] {e}\n{traceback.format_exc()}", file=sys.stderr)

    return {
        "applied": applied,
        "comments": comments,
        "skipped": 0,
        "section_actions": section_actions,
        "section_positions": section_positions,
    }


def extract_text(docx_path: str) -> str:
    """Extract full text from a .docx file."""
    doc = Document(docx_path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_parts = []
            for cell in row.cells:
                if cell.text.strip():
                    row_parts.append(cell.text.strip())
            if row_parts:
                parts.append(" | ".join(row_parts))
    return "\n\n".join(parts)


def extract_text_from_pdf(pdf_path: str) -> str:
    """Extract text from a PDF using pdfplumber."""
    import pdfplumber
    parts = []
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text and text.strip():
                parts.append(text.strip())
    return "\n\n".join(parts)


def create_docx_from_text(text: str, output_path: str) -> str:
    """
    Create a basic .docx from plain text (used for PDF → redlineable Word doc).
    Preserves paragraph structure.
    """
    doc = Document()
    for line in text.split('\n'):
        doc.add_paragraph(line)
    doc.save(output_path)
    return output_path
